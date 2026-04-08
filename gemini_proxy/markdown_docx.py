from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt

INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")
NUMBERED_LIST_RE = re.compile(r"^\d+\.\s+")
TABLE_SEPARATOR_RE = re.compile(r"^\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?$")


def render_markdown_to_docx(
    *,
    report_text: str,
    output_path: Path,
    title: str,
    meta_lines: list[tuple[str, str]],
    source_notes: list[str] | None = None,
) -> Path:
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)

    for label, value in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        _append_inline_runs(paragraph, value)

    if source_notes:
        paragraph = document.add_paragraph()
        paragraph.add_run("Источники: ").bold = True
        _append_inline_runs(paragraph, "; ".join(source_notes))

    _render_markdown_body(document, report_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _render_markdown_body(document: Document, report_text: str) -> None:
    lines = report_text.splitlines()
    paragraph_buffer: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        _append_inline_runs(paragraph, " ".join(part.strip() for part in paragraph_buffer if part.strip()))
        paragraph_buffer.clear()

    while index < len(lines):
        raw_line = lines[index].rstrip()
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            index += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            index = _render_code_block(document, lines, index)
            continue

        if line.startswith("# "):
            flush_paragraph()
            index += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            document.add_heading(line[3:].strip(), level=1)
            index += 1
            continue

        if line.startswith("### "):
            flush_paragraph()
            document.add_heading(line[4:].strip(), level=2)
            index += 1
            continue

        if line.startswith("#### "):
            flush_paragraph()
            document.add_heading(line[5:].strip(), level=3)
            index += 1
            continue

        if line in {"---", "***"}:
            flush_paragraph()
            document.add_paragraph()
            index += 1
            continue

        if _is_table_start(lines, index):
            flush_paragraph()
            index = _render_table(document, lines, index)
            continue

        if NUMBERED_LIST_RE.match(line):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            _append_inline_runs(paragraph, NUMBERED_LIST_RE.sub("", line, count=1).strip())
            index += 1
            continue

        if line.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_runs(paragraph, line[2:].strip())
            index += 1
            continue

        if line.startswith("> "):
            flush_paragraph()
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[2:].strip())
            run.italic = True
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def _render_code_block(document: Document, lines: list[str], start_index: int) -> int:
    language = lines[start_index].strip()[3:].strip()
    code_lines: list[str] = []
    index = start_index + 1
    while index < len(lines):
        current = lines[index].rstrip("\n")
        if current.strip().startswith("```"):
            break
        code_lines.append(current)
        index += 1

    paragraph = document.add_paragraph()
    if language:
        paragraph.add_run(f"{language}\n").bold = True

    for code_index, code_line in enumerate(code_lines):
        run = paragraph.add_run(code_line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if code_index < len(code_lines) - 1:
            run.add_break()

    return index + 1


def _is_table_start(lines: list[str], start_index: int) -> bool:
    if start_index + 1 >= len(lines):
        return False
    return "|" in lines[start_index] and bool(TABLE_SEPARATOR_RE.match(lines[start_index + 1].strip()))


def _render_table(document: Document, lines: list[str], start_index: int) -> int:
    table_lines = [lines[start_index], lines[start_index + 1]]
    index = start_index + 2
    while index < len(lines):
        candidate = lines[index].rstrip()
        if not candidate.strip() or "|" not in candidate:
            break
        table_lines.append(candidate)
        index += 1

    rows = [_split_table_row(line) for line in table_lines if line.strip()]
    if len(rows) < 2:
        return index

    header = rows[0]
    data_rows = [row for row in rows[2:] if row]
    column_count = len(header)
    for row in data_rows:
        column_count = max(column_count, len(row))

    table = document.add_table(rows=1 + len(data_rows), cols=column_count)
    table.style = "Table Grid"

    for column_index in range(column_count):
        cell = table.rows[0].cells[column_index]
        _append_inline_runs(cell.paragraphs[0], header[column_index] if column_index < len(header) else "")
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row_index, row in enumerate(data_rows, start=1):
        for column_index in range(column_count):
            cell = table.rows[row_index].cells[column_index]
            _append_inline_runs(cell.paragraphs[0], row[column_index] if column_index < len(row) else "")

    return index


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _append_inline_runs(paragraph, text: str) -> None:
    for token in INLINE_TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            continue
        if token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            continue
        if token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        if token.startswith("[") and "](" in token and token.endswith(")"):
            label, url = token[1:-1].split("](", maxsplit=1)
            run = paragraph.add_run(f"{label} ({url})")
            run.underline = True
            continue
        paragraph.add_run(token)
